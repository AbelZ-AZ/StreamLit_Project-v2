import streamlit as st
import numpy as np
import pandas as pd
import json
import urllib.request
import re
import base64
import os
from datetime import datetime
from openai import OpenAI
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Optional libraries for document parsing
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


# ==============================================================================
# --- PAGE CONFIGURATION ---
# ==============================================================================
st.set_page_config(
    page_title="Price Reasonableness & SVP Assistant",
    page_icon="⚖️",
    layout="wide"
)


# ==============================================================================
# --- HELPER FUNCTIONS & KNOWLEDGE BASE LOADERS ---
# ==============================================================================
def find_template_file(filename):
    """Dynamically locates template files across possible repository paths."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    possible_paths = [
        os.path.join(base_dir, "templates", filename),
        os.path.join(base_dir, "StreamLit_Project", "templates", filename),
        os.path.join("templates", filename),
        os.path.join("StreamLit_Project", "templates", filename)
    ]
    for path in possible_paths:
        if os.path.exists(path):
            return path
    return None


@st.cache_data
def load_docx_from_repo(file_path):
    """Loads and extracts text from a .docx file stored in the GitHub repository."""
    if not file_path or not os.path.exists(file_path):
        return ""
    if docx is None:
        st.error("The 'python-docx' package is missing. Please run 'pip install python-docx'.")
        return ""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Error reading template from path `{file_path}`: {e}")
        return ""


@st.cache_data(ttl=3600)
def get_exchange_rates_to_sgd():
    """Fetches live exchange rates or uses static fallbacks."""
    fallback_rates = {
        "SGD": 1.0, "USD": 1.35, "CNY": 0.19, "RMB": 0.19,
        "JPY": 0.0088, "EUR": 1.46, "GBP": 1.72, "AUD": 0.88,
        "MYR": 0.30, "HKD": 0.17
    }
    try:
        url = "https://open.er-api.com/v6/latest/SGD"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=3) as res:
            if res.status == 200:
                data = json.loads(res.read().decode('utf-8'))
                rates = data.get("rates", {})
                if rates:
                    converted_rates = {"SGD": 1.0}
                    for curr, rate in rates.items():
                        if rate > 0:
                            converted_rates[curr.upper()] = 1.0 / rate
                    return converted_rates
    except Exception:
        pass
    return fallback_rates


def convert_to_sgd(amount, currency_code, rates):
    """Converts a given price amount to SGD using exchange rate lookups."""
    curr = str(currency_code).upper().strip()
    rate = rates.get(curr)
    if rate is None:
        rate = rates.get("USD", 1.35) if curr != "SGD" else 1.0
    return round(float(amount) * rate, 2)


# ==============================================================================
# --- RAG RETRIEVAL ENGINE ---
# ==============================================================================
def chunk_text(text, chunk_size=120, overlap=25):
    if not text:
        return []
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    return chunks


def retrieve_svp_guidelines(query, text_chunks, top_k=2):
    if not text_chunks:
        return "No SVP policy document found in repository templates."
    vectorizer = TfidfVectorizer()
    corpus = text_chunks + [query]
    tfidf_matrix = vectorizer.fit_transform(corpus)
    scores = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1]).flatten()
    top_indices = scores.argsort()[-top_k:][::-1]
    return "\n".join([text_chunks[i] for i in top_indices if scores[i] > 0.02])


# ==============================================================================
# --- FILE & MULTIMODAL VISION OCR HELPERS ---
# ==============================================================================
def extract_text_from_file(uploaded_file):
    file_ext = uploaded_file.name.split(".")[-1].lower()
    extracted_text = ""
    try:
        if file_ext == "pdf" and pypdf:
            reader = pypdf.PdfReader(uploaded_file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
        elif file_ext in ["docx", "doc"] and docx:
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    extracted_text += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
    except Exception as e:
        st.warning(f"Note: Error extracting text: {e}")
    return extracted_text


def parse_quotation_with_llm(file_text, api_key):
    if not file_text.strip():
        return []
    try:
        client = OpenAI(api_key=api_key)
        prompt = (
            "Extract line items from this quotation text into JSON format:\n"
            '{"items": [{"Item Description": "desc", "Quantity": 1, "Unit Rate (SGD, incl. GST)": 100.00}]}\n'
            f"DOCUMENT TEXT:\n{file_text[:4000]}"
        )
        response = client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0
        )
        data = json.loads(response.choices[0].message.content.strip())
        return data.get("items", [])
    except Exception:
        return []


def parse_image_with_vision(uploaded_file, api_key):
    try:
        client = OpenAI(api_key=api_key)
        bytes_data = uploaded_file.getvalue()
        base64_image = base64.b64encode(bytes_data).decode("utf-8")
        
        prompt = (
            "Analyze this quotation or e-commerce screenshot and extract product line items into JSON.\n"
            'Format: {"items": [{"Item Description": "Name", "Quantity": 1, "Unit Rate (SGD, incl. GST)": 99.00}]}'
        )
        
        response = client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }
            ],
            temperature=0.0
        )
        data = json.loads(response.choices[0].message.content.strip())
        return data.get("items", [])
    except Exception as e:
        st.error(f"Error reading screenshot: {e}")
        return []


def scan_and_extract_file(uploaded_file, api_key):
    file_ext = uploaded_file.name.split(".")[-1].lower()
    parsed_items = []
    
    if file_ext in ["xlsx", "xls"]:
        try:
            df_uploaded = pd.read_excel(uploaded_file)
            df_uploaded.columns = df_uploaded.columns.str.strip().str.lower()
            for _, row in df_uploaded.iterrows():
                desc = str(row.get("item description", row.get("description", row.get("item", "")))).strip()
                qty = float(row.get("quantity", row.get("qty", 1)))
                rate = float(row.get("unit rate (sgd, incl. gst)", row.get("unit rate", row.get("unit price", row.get("price", 0.0)))))
                if desc and rate > 0:
                    parsed_items.append({"Item Description": desc, "Quantity": int(qty), "Unit Rate (SGD, incl. GST)": round(rate, 2)})
        except Exception as e:
            st.error(f"Error parsing Excel: {e}")
    elif file_ext in ["png", "jpg", "jpeg"]:
        if not api_key:
            st.error("🔑 OpenAI API Key is required for Vision Screenshot parsing.")
            return []
        parsed_items = parse_image_with_vision(uploaded_file, api_key)
    elif file_ext in ["pdf", "docx", "doc"]:
        if not api_key:
            st.error("🔑 OpenAI API Key required for document scanning.")
            return []
        raw_text = extract_text_from_file(uploaded_file)
        if raw_text:
            parsed_items = parse_quotation_with_llm(raw_text, api_key)
        else:
            st.error("Could not extract readable text from document.")
    return parsed_items


# ==============================================================================
# --- MARKET BENCHMARKING ENGINE (NATIVE OPENAI RESPONSES SEARCH) ---
# ==============================================================================
def search_market_prices_openai(item_description, user_openai_key, retrieved_rules, exchange_rates):
    """
    Executes live web search using OpenAI Responses API,
    prioritizing major Singapore retailers & local SME suppliers, explicitly excluding second-hand C2C platforms like Carousell.
    """
    client = OpenAI(api_key=user_openai_key)

    system_instructions = (
        "You are a procurement analysis bot specializing in Singapore market benchmarking & SVP Policy Compliance.\n"
        f"--- RETRIEVED REPO SVP POLICY RULES ---\n{retrieved_rules}\n---------------------\n"
        "Your task:\n"
        "1. Search live online sources for the specified product, prioritizing MAJOR Singapore authorized retailers and official distributors (e.g., Shopee SG, Lazada SG, Amazon SG, Challenger, Courts, Gain City, Harvey Norman SG) as well as established local SME suppliers/specialist vendors.\n"
        "2. EXCLUDE CAROUSELL & SECOND-HAND C2C PLATFORMS: Strictly DO NOT include listings from peer-to-peer or second-hand marketplace platforms such as Carousell, eBay second-hand listings, or personal C2C seller apps. Retrieve ONLY brand-new retail items from commercial suppliers.\n"
        "3. Aim to collect 10 distinct pricing data points from official Singapore sources first. Supplement with reputable overseas sources (e.g., official Amazon US, Taobao brand stores) only if fewer than 10 local Singapore price points are available.\n"
        "4. ACCURATE PRICE EXTRACTION: Carefully extract the exact numerical value of the final nett unit price (inclusive of GST/taxes where applicable). Ensure the price is returned as a pure numeric number (e.g., 149.90, not string text like 'S$149.90' or '149.90 SGD').\n"
        "5. Explicitly specify the original unit price and 3-letter currency code (e.g., SGD, USD, CNY, JPY) for each item. Original unit price MUST equal the exact extracted nett price.\n"
        "6. Include direct URLs or domain links for source verification.\n"
        "7. Provide a 'suggestion_action' advising if the price is fair or requires specific justifications under SVP policy.\n"
        "8. Return your result strictly in raw JSON format without markdown code blocks, adhering to this structure:\n"
        "{\n"
        '  "prices_found": [\n'
        '    {"source_name": "Lazada SG (Official Store)", "original_price": 129.00, "currency": "SGD", "region": "Singapore", "url": "https://www.lazada.sg/..."},\n'
        '    {"source_name": "Challenger SG", "original_price": 135.00, "currency": "SGD", "region": "Singapore", "url": "https://www.hachi.tech/..."}\n'
        "  ],\n"
        '  "suggestion_action": "Suggested procurement steps..."\n'
        "}"
    )

    user_prompt = f"Find current live brand-new nett market unit prices (targeting 10 sources, prioritizing major SG retailers and local SG suppliers, excluding Carousell) for: {item_description}"

    response = client.responses.create(
        model="gpt-4o",
        tools=[{"type": "web_search"}],
        input=[
            {"role": "system", "content": system_instructions},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.2
    )

    raw_text = response.output_text.strip()
    
    if raw_text.startswith("```json"):
        raw_text = raw_text[7:]
    elif raw_text.startswith("```"):
        raw_text = raw_text[3:]

    if raw_text.endswith("```"):
        raw_text = raw_text[:-3]

    data = json.loads(raw_text.strip())
    web_items = data.get("prices_found", [])

    for item in web_items:
        try:
            orig_p = float(re.sub(r"[^\d\.]", "", str(item.get("original_price", 0.0))))
        except ValueError:
            orig_p = 0.0

        item["original_price"] = orig_p
        curr = str(item.get("currency", "SGD")).upper().strip()
        item["currency"] = curr
        item["price_sgd"] = orig_p if curr == "SGD" else convert_to_sgd(orig_p, curr, exchange_rates)

    return web_items, data.get("suggestion_action", "")


# ==============================================================================
# --- INITIALIZE SESSION STATE ---
# ==============================================================================
if "search_history" not in st.session_state:
    st.session_state.search_history = []
if "assessment_results" not in st.session_state:
    st.session_state.assessment_results = []
if "assessed" not in st.session_state:
    st.session_state.assessed = False
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = [
        {"role": "assistant", "content": "Hello! I am your SVP Policy Guide. Your organization's SVP Policy document is pre-loaded from GitHub. Ask me anything about rules or compliance!"}
    ]

exchange_rates = get_exchange_rates_to_sgd()

# --- LOAD REPOSITORY KNOWLEDGE BASE AUTOMATICALLY ---
POLICY_PATH = find_template_file("svp_process.docx")
svp_policy_text = load_docx_from_repo(POLICY_PATH) if POLICY_PATH else ""
rag_chunks = chunk_text(svp_policy_text)

TEMPLATE_PATH = find_template_file("email_template.docx")
email_template_text = load_docx_from_repo(TEMPLATE_PATH) if TEMPLATE_PATH else ""


# ==============================================================================
# --- SIDEBAR: CONFIGURATION & KNOWLEDGE BASE STATUS ---
# ==============================================================================
st.sidebar.title("⚙️ Configuration")
user_api_key = st.sidebar.text_input(
    label="OpenAI API Key",
    type="password",
    placeholder="sk-...",
    help="Enter your OpenAI API key to run live web benchmarking & policy search."
)

st.sidebar.divider()
st.sidebar.title("📚 Repo Knowledge Base")

if svp_policy_text:
    st.sidebar.success("✅ SVP Policy Doc Loaded (`svp_process.docx`)")
else:
    st.sidebar.error("❌ `svp_process.docx` not found in `templates/` folder.")

if email_template_text:
    st.sidebar.success("✅ Email Template Loaded (`email_template.docx`)")
else:
    st.sidebar.error("❌ `email_template.docx` not found in `templates/` folder.")

st.sidebar.divider()
st.sidebar.title("📜 Search History")
if st.session_state.search_history:
    st.sidebar.caption(f"Total Assessments Run: **{len(st.session_state.search_history)}**")
    for idx, hist in enumerate(reversed(st.session_state.search_history)):
        sup = hist.get("supplier_name") or "Unknown Supplier"
        ref = hist.get("quotation_ref") or "No Ref"
        ts = hist.get("timestamp", "")
        cost = hist.get("total_cost", 0.0)
        
        with st.sidebar.expander(f"🕒 {ts} - {sup[:15]}...", expanded=False):
            st.markdown(f"**Supplier:** {sup}")
            st.markdown(f"**Quote Ref:** {ref}")
            st.markdown(f"**Total Cost:** S${cost:,.2f}")
            st.markdown(f"**Items:** {len(hist.get('results', []))}")
            
    if st.sidebar.button("🗑️ Clear History", key="clear_hist_btn"):
        st.session_state.search_history = []
        st.rerun()
else:
    st.sidebar.caption("No assessment history yet.")


# ==============================================================================
# --- NAVIGATION TABS ---
# ==============================================================================
tab_assessment, tab_chat = st.tabs(["📊 Price Assessment & Approval", "💬 SVP Policy Q&A Chatbot"])


# ==============================================================================
# TAB 1: PRICE REASONABLENESS ASSESSMENT
# ==============================================================================
with tab_assessment:
    st.title("⚖️ Price Reasonableness Assessment & Approval Helper")
    st.subheader("Evaluate Small Value Purchases (SVP) via Singapore Sourcing, RAG Rules & IQR Analysis")
    st.divider()

    st.markdown("### 📄 Step 1: Upload Quotation & Supplier Details")
    col_sup1, col_sup2 = st.columns([2, 2])
    with col_sup1:
        supplier_name = st.text_input(label="Supplier Name", placeholder="e.g., Tech Supplies Pte Ltd")
    with col_sup2:
        quotation_ref = st.text_input(label="Quotation Reference / No.", placeholder="e.g., QUO-2026-0891")

    uploaded_quote = st.file_uploader(
        label="Upload Supplier Quotation or Screenshot (PDF, Excel, Word, PNG, JPG)",
        type=["pdf", "xlsx", "xls", "docx", "doc", "png", "jpg", "jpeg"],
        help="Upload a quote document or screenshot to scan items into Step 2."
    )

    if "line_items_data" not in st.session_state:
        st.session_state.line_items_data = pd.DataFrame([
            {"Item Description": "Logitech MX Master 3S Wireless Mouse", "Quantity": 1, "Unit Rate (SGD, incl. GST)": 139.00}
        ])

    scan_col1, _ = st.columns([1, 3])
    with scan_col1:
        scan_button = st.button("🔍 Scan & Extract Quotation Items", use_container_width=True)

    if scan_button:
        if uploaded_quote is None:
            st.warning("⚠️ Please upload a quotation file first before clicking scan.")
        else:
            with st.spinner(f"Scanning `{uploaded_quote.name}` and extracting line item details..."):
                extracted = scan_and_extract_file(uploaded_quote, user_api_key)
                if extracted:
                    st.session_state.line_items_data = pd.DataFrame(extracted)
                    st.success(f"✅ Successfully extracted {len(extracted)} line item(s) into Step 2 below!")
                else:
                    st.warning("Could not automatically extract line items. Please enter details manually in Step 2.")

    st.divider()

    st.markdown("### 📝 Step 2: Line Items & Pricing Breakdown")
    st.caption("Populated from scanned quote file or entered manually below. Total costs compute automatically.")

    edited_df = st.data_editor(
        st.session_state.line_items_data
